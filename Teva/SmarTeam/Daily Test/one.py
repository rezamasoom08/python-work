import tkinter as tk
from tkinter import *
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.ie.service import Service as IEService
from selenium.webdriver.ie.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
from PIL import Image
from docx import Document
from docx.shared import Inches
import time
import os
from datetime import datetime

current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
filename_datetime = datetime.now().strftime("%Y-%m-%d")

def on_submit():
    driver = webdriver.Edge()
    driver.maximize_window()
    
    username = entry_username.get()
    password = entry_password.get()
    
    screenshots = []

    urls = ['https://ilstweb-st.teva.corp/WebEditor/Authentication/LoginPage.aspx?ReturnUrl=%2fwebeditor',
        'https://ilstweb03.teva.corp/WebEditor/Authentication/LoginPage.aspx?ReturnUrl=%2fWebEditor',
        'https://ilstweb04.teva.corp/WebEditor/Authentication/LoginPage.aspx?ReturnUrl=%2fWebEditor',
        'https://ilstweb05.teva.corp/WebEditor/Authentication/LoginPage.aspx?ReturnUrl=%2fWebEditor',
        'https://ilstweb06.teva.corp/WebEditor/Authentication/LoginPage.aspx?ReturnUrl=%2fWebEditor',
        'https://ilstweb-view.teva.corp/SmartView/Authentication/LoginPage.aspx?ReturnUrl=%2fSmartView%2fSmartView.aspx',
        'https://ilstweb07.teva.corp/SmartView/Authentication/LoginPage.aspx?ReturnUrl=%2fSmartView%2fSmartView.aspx',
        'https://ilstweb08.teva.corp/SmartView/Authentication/LoginPage.aspx?ReturnUrl=%2fSmartView%2fSmartView.aspx']

    for url in urls:
        driver.get(url)
        
        try:
            user = driver.find_element("id", 'AuthenticatedLogin1_tbUsername')
            passwd = driver.find_element("id", 'AuthenticatedLogin1_tbPassword')
                        
            user.send_keys(username)
            passwd.send_keys(password + Keys.RETURN)
                        
            time.sleep(15)  # Wait for login to process, adjust as necessary

            screenshot_path = f'screenshot_{urls.index(url)}.png'
            driver.save_screenshot(screenshot_path)
            screenshots.append(screenshot_path)
        
        except Exception as e:
            print(f'Error with URL {url}: {e}')
    
    driver.quit()
    create_word_document(screenshots)

def create_word_document(screenshots):
    doc = Document()
    doc.add_heading('Smart Web & View', level=1)
    doc.add_paragraph(f"Date and Time: {current_datetime}")

    for screenshot in screenshots:
        doc.add_picture(screenshot, width=Inches(6))
    
    doc.save(f'C:/Smarteam/Smart Web & View {filename_datetime}.docx')
    
    for screenshot in screenshots:
        os.remove(screenshot)

def dummy():
    driver = webdriver.Edge()
    driver.maximize_window()
        
    username = entry_username.get()
    password = entry_password.get()
    
    screenshots = []

    urls = ['https://ilstweb-st.teva.corp/WebEditor/Authentication/LoginPage.aspx?ReturnUrl=%2fwebeditor']
    
    for url in urls:
        driver.get(url)
    
        try:
            
            user = driver.find_element("id", 'AuthenticatedLogin1_tbUsername')
            passwd = driver.find_element("id", 'AuthenticatedLogin1_tbPassword')
                        
            user.send_keys(username)
            passwd.send_keys(password + Keys.RETURN)
                        
            time.sleep(15) # Wait for login to process, adjust as necessary

            screenshot_path = f'screenshot_{urls.index(url)}.png'
            driver.save_screenshot(screenshot_path)
            screenshots.append(screenshot_path)

        except Exception as e:
            print(f'Error with URL {url}: {e}')

    driver.quit()

# Create a Word document and add screenshots
    doc = Document()
    doc.add_heading('Dummy User', level=1)
    doc.add_paragraph(f"Captured at: {current_datetime}")

    for screenshot in screenshots:
        doc.add_picture(screenshot, width=Inches(6))
        

# Save the document
    doc.save(f'C:/Smarteam/Dummy User {filename_datetime}.docx')

# Cleanup screenshots
    for screenshot in screenshots:
        os.remove(screenshot)

def portal():
    driver = webdriver.Edge()  # or webdriver.Firefox(), etc.
    driver.maximize_window()

# Open the web page
    portal_url = 'http://gaus.hr.teva.corp/queryPortal/smarteam/'  # replace with the desired URL
    driver.get(portal_url)

    # Wait for the page to load completely
    time.sleep(5)  # adjust as needed

    # Perform actions and take screenshots
    actions = ActionChains(driver)
    screenshots = []
    for i in range(1):  # Adjust the range for the number of screenshots needed

    # Perform some actions like clicking
        element = driver.find_element(By.ID, "btn2")  # Replace with the appropriate element
        actions.move_to_element(element).click().perform()
        time.sleep(5)  # Adjust as needed for actions to complete

    # Take a screenshot
        screenshot_filename = f"screenshot_{i+1}.png"
        driver.save_screenshot(screenshot_filename)
        screenshots.append(screenshot_filename)

        element = driver.find_element(By.ID, "btn3")  # Replace with the appropriate element
        actions.move_to_element(element).click().perform()
        time.sleep(5)  # Adjust as needed for actions to complete

    # Take a screenshot
        screenshot_filename = f"screenshot_{i+2}.png"
        driver.save_screenshot(screenshot_filename)
        screenshots.append(screenshot_filename)

    driver.quit()

# Create a Word document and add screenshots
    doc = Document()
    doc.add_heading('Query Portal', level=1)
    doc.add_paragraph(f"Captured at: {current_datetime}")

    for screenshot in screenshots:
        doc.add_picture(screenshot, width=Inches(6))
        

# Save the document
    doc.save(f'C:/Smarteam/Query_Portal {filename_datetime}.docx')

# Cleanup screenshots
    for screenshot in screenshots:
        os.remove(screenshot)

    os.startfile('C:/Smarteam/')

root = tk.Tk()
root.title('SMARTEAM SCREENSHOT ASSISTANT')
root.maxsize(400, 300)
root.minsize(400, 300)
root.configure(background='grey')

tk.Label(root, text='SMARTEAM LOGIN', bg='grey', font=('Ubuntu Mono', 12)).grid(row=0, column=0, columnspan=2, pady=15, sticky='n')
tk.Label(root, text='Username', bg='grey', font=('Ubuntu Mono', 10)).grid(row=1, column=0, padx=5, pady=3, ipadx=10, sticky='e')
tk.Label(root, text='Password', bg='grey', font=('Ubuntu Mono', 10)).grid(row=2, column=0, padx=5, pady=3, ipadx=10, sticky='e')

entry_username = tk.Entry(root, font=('Ubuntu Mono', 10), width=25)
entry_password = tk.Entry(root, show='*', font=('Ubuntu Mono', 10), width=25)

entry_username.grid(row=1, column=1, ipadx=50, pady=3)
entry_password.grid(row=2, column=1, ipadx=50, pady=3)

tk.Button(root, text='Web & View', fg='Black', bg='grey', width=20, font=('Ubuntu Mono', 10), command=on_submit).grid(row=4, column=0, columnspan=2, pady=12, sticky='n')
tk.Button(root, text='Dummy User', fg='Black', width=20, command=dummy, bg='grey', font=('Ubuntu Mono', 10)).grid(row=5, column=0, columnspan=2, pady=12, sticky='n')
tk.Button(root, text='Query Portal', fg='Black', command=portal, width=20, bg='grey', font=('Ubuntu Mono', 10)).grid(row=6, column=0, columnspan=2, pady=10, sticky='n')

status_bar = tk.Label(root, text='This App is designed and created by MAX. Copyright @2024', bg='grey', justify=tk.CENTER, font=("Ubuntu Mono", 8))
status_bar.grid(row=7, column=0, columnspan=3, pady=15, sticky='s')

root.mainloop()
